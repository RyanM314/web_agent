"""
app/services/document_parser.py
────────────────────────────────
Stateless functions that extract plain text from uploaded files.

Supported formats
─────────────────
• TXT  — decoded with chardet-assisted charset detection
• PDF  — text layer extracted with pypdf; falls back to page-by-page extraction
• DOCX — paragraph + table text extracted with python-docx

All public functions return a plain string.
Failures raise DocumentParseError so callers can produce a clean 422 response.
"""

from __future__ import annotations

import logging
from pathlib import Path

import chardet

from app.core.exceptions import DocumentParseError

logger = logging.getLogger(__name__)


# ── Dispatcher ─────────────────────────────────────────────────────────────────

def parse_document(file_path: Path) -> str:
    """
    Extract and return the full text content of a document file.

    Parameters
    ----------
    file_path : Path
        Absolute (or cwd-relative) path to the saved upload.

    Returns
    -------
    str
        Extracted text, with whitespace normalised.

    Raises
    ------
    DocumentParseError
        If the file cannot be read or the format is unsupported.
    """
    ext = file_path.suffix.lower().lstrip(".")
    parsers = {"txt": _parse_txt, "pdf": _parse_pdf, "docx": _parse_docx}

    parser = parsers.get(ext)
    if parser is None:
        raise DocumentParseError(
            f"No parser available for '.{ext}' files. "
            f"Supported: {', '.join(parsers)}."
        )

    try:
        text = parser(file_path)
    except DocumentParseError:
        raise
    except Exception as exc:  # pragma: no cover
        raise DocumentParseError(
            f"Failed to parse '{file_path.name}': {exc}"
        ) from exc

    text = _normalise(text)
    if not text:
        raise DocumentParseError(
            f"'{file_path.name}' was parsed successfully but contains no readable text."
        )

    logger.info("Parsed '%s' → %d characters.", file_path.name, len(text))
    return text


# ── Individual parsers ─────────────────────────────────────────────────────────

def _parse_txt(path: Path) -> str:
    """Read a plain-text file with automatic charset detection."""
    raw = path.read_bytes()
    detection = chardet.detect(raw)
    encoding = detection.get("encoding") or "utf-8"
    confidence = detection.get("confidence", 0)

    logger.debug(
        "TXT charset detection: encoding=%s confidence=%.2f", encoding, confidence
    )

    try:
        return raw.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        # Last-resort: utf-8 with replacement characters
        logger.warning(
            "Could not decode '%s' as %s; falling back to utf-8 with replacement.",
            path.name,
            encoding,
        )
        return raw.decode("utf-8", errors="replace")


def _parse_pdf(path: Path) -> str:
    """Extract the text layer from a PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise DocumentParseError("pypdf is not installed.") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []

    for page_num, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            logger.warning("Could not extract page %d: %s", page_num, exc)
            page_text = ""
        pages.append(page_text)

    combined = "\n".join(pages)
    logger.debug("PDF '%s': %d pages, %d chars extracted.", path.name, len(pages), len(combined))
    return combined


def _parse_docx(path: Path) -> str:
    """Extract text from all paragraphs and tables in a DOCX file."""
    try:
        from docx import Document  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise DocumentParseError("python-docx is not installed.") from exc

    doc = Document(str(path))
    parts: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _normalise(text: str) -> str:
    """
    Collapse runs of blank lines and strip leading/trailing whitespace.
    Preserves single blank lines (paragraph breaks) for readability.
    """
    import re
    # Replace 3+ consecutive newlines with exactly two
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
