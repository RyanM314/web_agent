"""
app/services/export_service.py
──────────────────────────────
FR4 — Output Generation Export

Converts generated markdown notes into downloadable file formats:
• PDF (via fpdf2)
• DOCX (via python-docx)

Both functions accept raw markdown text and produce an in-memory byte stream
suitable for a streaming HTTP response.

Table Support
─────────────
Markdown pipe tables (| col1 | col2 |) are parsed and rendered as proper
tables in both PDF and DOCX outputs.

Original Document Preservation
──────────────────────────────
When an original document path is provided, tables from the original are
extracted and included in the exported document to preserve structural
elements that cannot be represented in plain markdown.
"""

from __future__ import annotations

import io
import logging
import os
import re
from pathlib import Path
from typing import Any

from app.core.exceptions import ExportError

logger = logging.getLogger(__name__)


# ── Public API ──────────────────────────────────────────────────────────────────

def export_to_pdf(
    markdown_text: str,
    title: str = "Updated Notes",
    original_document_path: str | None = None,
) -> io.BytesIO:
    """
    Convert markdown text to a PDF document.

    Parameters
    ----------
    markdown_text : The notes content (may contain markdown formatting).
    title         : Document title displayed on the first page.
    original_document_path : Optional path to original file for table extraction.

    Returns
    -------
    io.BytesIO — in-memory PDF file ready for download.

    Raises
    ------
    ExportError if PDF generation fails.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise ExportError("fpdf2 is not installed. Run: pip install fpdf2") from None

    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=20)

        # Add a Unicode-capable font for special characters
        _add_unicode_font(pdf)

        _add_pdf_title_page(pdf, title)

        # Extract tables from original document if provided
        original_tables: list[list[list[str]]] = []
        if original_document_path and os.path.exists(original_document_path):
            try:
                original_tables = _extract_tables_from_document(original_document_path)
            except Exception as exc:
                logger.debug("Could not extract tables from original document: %s", exc)

        # Parse markdown into sections and render
        _render_pdf_content(pdf, markdown_text, original_tables)

        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)
        logger.info("PDF generated: %d bytes, title='%s'.", buf.getbuffer().nbytes, title)
        return buf

    except Exception as exc:
        raise ExportError(f"PDF generation failed: {exc}") from exc


def export_to_docx(
    markdown_text: str,
    title: str = "Updated Notes",
    original_document_path: str | None = None,
) -> io.BytesIO:
    """
    Convert markdown text to a DOCX document.

    Parameters
    ----------
    markdown_text : The notes content (may contain markdown formatting).
    title         : Document title displayed at the top.
    original_document_path : Optional path to original file for table extraction.

    Returns
    -------
    io.BytesIO — in-memory DOCX file ready for download.

    Raises
    ------
    ExportError if DOCX generation fails.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ExportError("python-docx is not installed. It should be in requirements.txt.") from None

    try:
        doc = Document()

        # ── Title ───────────────────────────────────────────────────────────────
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle with timestamp
        from datetime import datetime
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # spacing

        # Extract tables from original document if provided
        original_tables: list[list[list[str]]] = []
        original_has_images: bool = False
        if original_document_path and os.path.exists(original_document_path):
            try:
                original_tables = _extract_tables_from_document(original_document_path)
                original_has_images = _detect_images_in_document(original_document_path)
            except Exception as exc:
                logger.debug("Could not extract tables from original document: %s", exc)

        # ── Parse and render markdown ──────────────────────────────────────────
        _render_docx_content(doc, markdown_text, original_tables, original_has_images)

        # Add preserved tables from original at the end
        if original_tables:
            doc.add_page_break()
            doc.add_heading("Preserved Tables from Original Document", level=1)
            for table_idx, table_data in enumerate(original_tables, start=1):
                doc.add_heading(f"Table {table_idx}", level=2)
                _add_docx_table(doc, table_data)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        logger.info("DOCX generated: %d bytes, title='%s'.", buf.getbuffer().nbytes, title)
        return buf

    except ExportError:
        raise
    except Exception as exc:
        raise ExportError(f"DOCX generation failed: {exc}") from exc


# ── PDF rendering ───────────────────────────────────────────────────────────────

def _render_pdf_content(pdf, markdown_text: str, original_tables: list[list[list[str]]]) -> None:
    """Render markdown content to PDF with table support."""
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Check for markdown table
        if _is_markdown_table_line(line):
            table_lines = []
            while i < len(lines) and _is_markdown_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1
                # Also collect separator lines (|---|---|)
                if i < len(lines) and re.match(r"^\s*[\-|:]+\s*$", lines[i]):
                    table_lines.append(lines[i])
                    i += 1
            _render_pdf_table(pdf, table_lines)
            continue

        _render_pdf_line(pdf, line)
        i += 1

    # Add preserved original tables at the end
    if original_tables:
        pdf.add_page()
        pdf.set_font(pdf.font_family, "B", 14)
        pdf.cell(0, 10, "Preserved Tables from Original Document", ln=True)
        pdf.ln(5)
        for table_idx, table_data in enumerate(original_tables, start=1):
            pdf.set_font(pdf.font_family, "B", 12)
            pdf.cell(0, 8, f"Table {table_idx}", ln=True)
            pdf.ln(2)
            _render_pdf_table_from_data(pdf, table_data)
            pdf.ln(5)


def _is_markdown_table_line(line: str) -> bool:
    """Check if a line is part of a markdown table."""
    return "|" in line and line.strip().startswith("|")


def _render_pdf_table(pdf, table_lines: list[str]) -> None:
    """Render a markdown pipe table in PDF."""
    if not table_lines:
        return

    # Parse table rows (skip separator lines)
    rows: list[list[str]] = []
    for line in table_lines:
        if re.match(r"^\s*[\-|:]+\s*$", line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    _render_pdf_table_from_data(pdf, rows)


def _render_pdf_table_from_data(pdf, rows: list[list[str]]) -> None:
    """Render a 2D list of cells as a table in PDF."""
    if not rows:
        return

    # Calculate column widths
    num_cols = max(len(row) for row in rows)
    col_widths = [0] * num_cols
    for row in rows:
        for i, cell in enumerate(row):
            if i < num_cols:
                # Approximate width based on content length
                col_widths[i] = max(col_widths[i], min(len(cell) * 2.5 + 4, 50))

    # Normalize to page width
    total_width = sum(col_widths)
    available_width = pdf.epw
    if total_width > available_width:
        scale = available_width / total_width
        col_widths = [w * scale for w in col_widths]

    font_name = pdf.font_family
    line_height = 6

    pdf.set_font(font_name, "", 10)

    for row_idx, row in enumerate(rows):
        # Pad row to match column count
        padded_row = row + [""] * (num_cols - len(row))

        if row_idx == 0:
            pdf.set_font(font_name, "B", 10)
        else:
            pdf.set_font(font_name, "", 10)

        for col_idx, cell in enumerate(padded_row):
            if col_idx < num_cols:
                x = pdf.get_x()
                y = pdf.get_y()
                w = col_widths[col_idx]

                # Save current position
                pdf.set_xy(x, y)
                pdf.multi_cell(w, line_height, cell, border=1)

                # Move to next column
                pdf.set_xy(x + w, y)

        pdf.ln(line_height)

    pdf.ln(3)


def _render_pdf_line(pdf, line: str) -> None:
    """Render a single markdown line into the PDF."""
    if not line.strip():
        pdf.ln(4)
        return

    font_name = pdf.font_family

    # Heading
    heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
    if heading_match:
        level = len(heading_match.group(1))
        text = _strip_markdown_inline(heading_match.group(2))
        sizes = {1: 18, 2: 16, 3: 14, 4: 12, 5: 11, 6: 11}
        pdf.set_font(font_name, "B", sizes.get(level, 12))
        pdf.ln(3)
        pdf.multi_cell(pdf.epw, 8, text)
        pdf.ln(2)
        return

    # Horizontal rule spanning full page width
    if re.match(r"^---+\s*$", line) or re.match(r"^\*\*\*+\s*$", line):
        pdf.set_draw_color(180, 180, 180)
        left = pdf.l_margin
        pdf.line(left, pdf.get_y(), left + pdf.epw, pdf.get_y())
        pdf.ln(4)
        return

    # Bullet list — use dash instead of bullet char for ASCII safety
    list_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
    if list_match:
        text = _strip_markdown_inline(list_match.group(2))
        pdf.set_font(font_name, "", 11)
        indent = len(list_match.group(1)) + 5
        left = pdf.l_margin
        pdf.set_x(left + indent)
        pdf.cell(6, 6, "- ")  # dash prefix, safe for all encodings
        pdf.multi_cell(0, 6, text)
        return

    # Numbered list
    num_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
    if num_match:
        text = _strip_markdown_inline(num_match.group(2))
        pdf.set_font(font_name, "", 11)
        indent = len(list_match.group(1)) + 5
        left = pdf.l_margin
        pdf.set_x(left + indent)
        num_prefix = num_match.group(0).split(".")[0] + "."
        pdf.cell(8, 6, num_prefix)
        pdf.multi_cell(0, 6, text)
        return

    # Code block delimiter — skip
    if line.strip().startswith("```"):
        return

    # Regular paragraph — use effective page width to prevent clipping
    text = _strip_markdown_inline(line)
    pdf.set_font(font_name, "", 11)
    pdf.multi_cell(pdf.epw, 6, text)


# ── DOCX rendering ──────────────────────────────────────────────────────────────

def _render_docx_content(
    doc,
    markdown_text: str,
    original_tables: list[list[list[str]]],
    original_has_images: bool,
) -> None:
    """Render markdown content to DOCX with table support."""
    lines = markdown_text.split("\n")
    in_code_block = False
    code_buffer: list[str] = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Check for markdown table
        if _is_markdown_table_line(line):
            table_lines = []
            while i < len(lines) and _is_markdown_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1
                # Also collect separator lines (|---|---|)
                if i < len(lines) and re.match(r"^\s*[\-|:]+\s*$", lines[i]):
                    table_lines.append(lines[i])
                    i += 1
            _add_docx_table_from_markdown(doc, table_lines)
            continue

        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                _add_docx_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = _strip_markdown_inline(heading_match.group(2))
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", line) or re.match(r"^\*\*\*+\s*$", line):
            _add_docx_horizontal_rule(doc)
            i += 1
            continue

        # Bullet list
        list_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if list_match:
            text = _strip_markdown_inline(list_match.group(2))
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if num_match:
            text = _strip_markdown_inline(num_match.group(2))
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Regular paragraph
        text = _strip_markdown_inline(line)
        doc.add_paragraph(text)
        i += 1


# ── Document table extraction ───────────────────────────────────────────────────

def _extract_tables_from_document(file_path: str) -> list[list[list[str]]]:
    """
    Extract tables from a document file (DOCX or PDF).
    Returns a list of tables, where each table is a 2D list of cell strings.
    """
    ext = Path(file_path).suffix.lower().lstrip(".")

    if ext == "docx":
        return _extract_tables_from_docx(file_path)
    elif ext == "pdf":
        return _extract_tables_from_pdf(file_path)
    else:
        return []


def _extract_tables_from_docx(file_path: str) -> list[list[list[str]]]:
    """Extract tables from a DOCX file."""
    try:
        from docx import Document as DocxReader
    except ImportError:
        return []

    doc = DocxReader(file_path)
    tables: list[list[list[str]]] = []

    for table in doc.tables:
        table_data: list[list[str]] = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            tables.append(table_data)

    return tables


def _extract_tables_from_pdf(file_path: str) -> list[list[list[str]]]:
    """Extract tables from a PDF file (best-effort using pypdf)."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return []

    reader = PdfReader(file_path)
    tables: list[list[list[str]]] = []

    for page in reader.pages:
        try:
            # pypdf has basic table extraction
            page_tables = page.extract_tables()
            for table in page_tables:
                cleaned = [[cell.strip() for cell in row if cell] for row in table if any(cell.strip() for cell in row)]
                if cleaned:
                    tables.append(cleaned)
        except Exception:
            # Table extraction may not work for all PDFs
            continue

    return tables


def _detect_images_in_document(file_path: str) -> bool:
    """Check if a DOCX document contains images."""
    ext = Path(file_path).suffix.lower().lstrip(".")

    if ext == "docx":
        try:
            from docx import Document as DocxReader
            doc = DocxReader(file_path)
            # Check for inline shapes (images)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    return True
        except Exception:
            pass

    elif ext == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            for page in reader.pages:
                if page.images:
                    return True
        except Exception:
            pass

    return False


# ── DOCX table helpers ──────────────────────────────────────────────────────────

def _add_docx_table_from_markdown(doc, table_lines: list[str]) -> None:
    """Convert markdown pipe table lines to a DOCX table."""
    from docx.shared import Inches, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    rows = []
    for line in table_lines:
        if re.match(r"^\s*[\-|:]+\s*$", line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    # Normalize row lengths
    max_cols = max(len(r) for r in rows)
    for row in rows:
        while len(row) < max_cols:
            row.append("")

    # Create table
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < max_cols:
                cell = row.cells[col_idx]
                cell.text = cell_text
                # Header row formatting
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table


def _add_docx_table(doc, table_data: list[list[str]]) -> None:
    """Add a pre-extracted table to a DOCX document."""
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if not table_data:
        return

    max_cols = max(len(r) for r in table_data)
    for row in table_data:
        while len(row) < max_cols:
            row.append("")

    table = doc.add_table(rows=len(table_data), cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < max_cols:
                cell = row.cells[col_idx]
                cell.text = cell_text
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    doc.add_paragraph()


# ── PDF internal helpers ────────────────────────────────────────────────────────

def _add_unicode_font(pdf) -> None:
    """
    Add a Unicode-compatible font for the PDF.

    Tries common TrueType fonts that support Unicode. Falls back to
    the built-in Helvetica (ASCII-only) if none are found.
    """
    # Common system font paths that support Unicode
    font_candidates = [
        # Windows TrueType fonts (DejaVu and Arial support Unicode)
        (r"C:\Windows\Fonts\DejaVuSans.ttf", r"C:\Windows\Fonts\DejaVuSans-Bold.ttf"),
        (r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
        (r"C:\Windows\Fonts\calibri.ttf", r"C:\Windows\Fonts\calibrib.ttf"),
        (r"C:\Windows\Fonts\times.ttf", r"C:\Windows\Fonts\timesbd.ttf"),
        # Linux DejaVu
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        # macOS
        ("/Library/Fonts/Arial.ttf", "/Library/Fonts/Arial Bold.ttf"),
    ]

    for regular_path, bold_path in font_candidates:
        if os.path.exists(regular_path):
            try:
                pdf.add_font("UniFont", "", regular_path, uni=True)
                if os.path.exists(bold_path):
                    pdf.add_font("UniFont", "B", bold_path, uni=True)
                pdf.set_font("UniFont", "", 11)
                logger.debug("PDF using Unicode font: %s", regular_path)
                return
            except Exception:
                continue

    # Fallback — use Helvetica (ASCII-only, but functional)
    logger.debug("PDF using fallback font: Helvetica")
    pdf.set_font("Helvetica", "", 11)


def _add_pdf_title_page(pdf, title: str) -> None:
    """Add a styled title page to the PDF."""
    pdf.add_page()
    font_name = pdf.font_family
    epw = pdf.epw  # effective page width (full width minus margins)
    left = pdf.l_margin
    pdf.set_font(font_name, "B", 24)
    pdf.ln(60)
    pdf.multi_cell(epw, 12, title, align="C")
    pdf.ln(10)
    pdf.set_font(font_name, "", 12)
    from datetime import datetime
    pdf.multi_cell(
        epw, 8, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", align="C"
    )
    pdf.ln(10)
    # Separator line spanning full page width
    pdf.set_draw_color(100, 100, 100)
    pdf.line(left, pdf.get_y(), left + epw, pdf.get_y())
    pdf.ln(10)


# ── DOCX internal helpers ────────────────────────────────────────────────────────

def _add_docx_code_block(doc, code_text: str) -> None:
    """Add a formatted code block to the DOCX."""
    from docx.shared import Inches, Pt, RGBColor

    for code_line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(code_line if code_line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(50, 50, 50)


def _add_docx_horizontal_rule(doc) -> None:
    """Add a horizontal rule line to the DOCX."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)


# ── Shared helpers ──────────────────────────────────────────────────────────────

def _strip_markdown_inline(text: str) -> str:
    """
    Remove markdown formatting characters from inline text.

    Handles: **bold**, *italic*, `inline code`, [links](url), ~~strikethrough~~
    """
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r" _(.+?)_", r"\1", text)
    text = re.sub(r"~~(.+?)~~", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text